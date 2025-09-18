import os
import pandas as pd
from PIL import Image
import json
import csv
import shutil
import fitz
from docx import Document
import pythoncom
import win32com.client
from pydub import AudioSegment
import wave
import contextlib
import cv2
import zipfile
import tarfile
import rarfile
import py7zr
import markdown
import html2text
import xml.etree.ElementTree as ET
import yaml
import h5py
import sqlite3
import base64
import matplotlib.pyplot as plt
from io import BytesIO
import tempfile

class ConverterController:
    def __init__(self):
        self.supported_formats = {
            'images': ['jpg', 'jpeg', 'png', 'bmp', 'gif', 'tiff', 'webp', 'ico', 'heic', 'psd', 'raw'],
            'data': ['csv', 'json', 'xlsx', 'xls', 'txt', 'xml', 'yaml', 'h5', 'sql'],
            'documents': ['pdf', 'doc', 'docx', 'txt', 'rtf', 'md', 'html', 'odt', 'epub'],
            'audio': ['mp3', 'wav', 'ogg', 'flac', 'm4a', 'wma', 'aac', 'aiff'],
            'video': ['mp4', 'avi', 'mov', 'wmv', 'flv', 'mkv', 'webm', 'mpeg', 'mpg'],
            'archives': ['zip', 'tar', 'gz', 'bz2', 'xz', '7z', 'rar'],
            'code': ['py', 'js', 'java', 'cpp', 'c', 'cs', 'rb', 'php', 'go', 'ts'],
            'other': ['bin', 'iso', 'dmg', 'db']
        }

        self.pil_format_mapping = {
            'jpg': 'JPEG', 'jpeg': 'JPEG', 'png': 'PNG', 'bmp': 'BMP',
            'gif': 'GIF', 'tiff': 'TIFF', 'webp': 'WEBP', 'ico': 'ICO',
            'heic': 'HEIF', 'raw': 'RAW'
        }

        self.conversion_stats = {
            'total': 0,
            'successful': 0,
            'failed': 0,
            'downloads_completed': 0,
            'by_category': {}
        }

    def get_supported_formats(self):
        all_formats = []
        for category in self.supported_formats.values():
            all_formats.extend(category)
        return sorted(all_formats)

    def get_output_formats_for_input(self, input_format):
        input_format = input_format.lower()
        category = None
        for cat, formats in self.supported_formats.items():
            if input_format in formats:
                category = cat
                break
        if not category:
            return []
        return [fmt.upper() for fmt in self.supported_formats[category] if fmt != input_format]

    def auto_convert_file(self, input_path, output_dir=None, preferred_format=None, options=None, progress_callback=None):
        if options is None:
            options = {}
        input_format = os.path.splitext(input_path)[1].lower().lstrip('.')
        possible_formats = self.get_output_formats_for_input(input_format)
        if not possible_formats:
            raise ValueError(f"No conversion formats available for: {input_format}")
        output_format = preferred_format.lower() if preferred_format and preferred_format.lower() in [f.lower() for f in possible_formats] else possible_formats[0].lower()
        if not output_dir:
            output_dir = os.path.dirname(input_path)
        output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + f".{output_format}")
        success = self.convert_file(input_path, output_path, output_format, options, progress_callback)
        self._update_stats(success)
        return output_path if success else None

    def convert_file(self, input_path, output_path, output_format, options, progress_callback=None):
        try:
            input_format = os.path.splitext(input_path)[1].lower().lstrip('.')
            output_format = output_format.lower()
            if progress_callback:
                progress_callback(25)
            if input_format in self.supported_formats['images'] and output_format in self.supported_formats['images']:
                success = self._convert_image(input_path, output_path, output_format, options)
            elif input_format in self.supported_formats['data'] and output_format in self.supported_formats['data']:
                success = self._convert_data(input_path, output_path, output_format)
            elif input_format in self.supported_formats['documents'] and output_format in self.supported_formats['documents']:
                success = self._convert_document(input_path, output_path, output_format)
            elif input_format in self.supported_formats['audio'] and output_format in self.supported_formats['audio']:
                success = self._convert_audio(input_path, output_path, output_format)
            elif input_format in self.supported_formats['video'] and output_format in self.supported_formats['video']:
                success = self._convert_video(input_path, output_path, output_format)
            elif input_format in self.supported_formats['archives'] and output_format in self.supported_formats['archives']:
                success = self._convert_archive(input_path, output_path, output_format)
            elif input_format in self.supported_formats['code'] and output_format in self.supported_formats['code']:
                success = self._copy_file(input_path, output_path)
            else:
                success = self._copy_file(input_path, output_path)
            if progress_callback:
                progress_callback(100 if success else 0)
            self._update_stats(success)
            if success:
                self.conversion_stats['downloads_completed'] += 1
            return success
        except Exception as e:
            self._update_stats(False)
            if progress_callback:
                progress_callback(0)
            raise e

    def convert_batch(self, input_dir, output_dir, output_format, options, progress_callback=None):
        try:
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            files = []
            for root, _, filenames in os.walk(input_dir):
                for filename in filenames:
                    files.append(os.path.join(root, filename))
            total_files = len(files)
            if total_files == 0:
                raise ValueError("No files found in the specified folder")
            success_count = 0
            for i, file_path in enumerate(files):
                try:
                    rel_path = os.path.relpath(file_path, input_dir)
                    output_path = os.path.join(output_dir, rel_path)
                    output_path = os.path.splitext(output_path)[0] + f'.{output_format}'
                    os.makedirs(os.path.dirname(output_path), exist_ok=True)
                    if self.convert_file(file_path, output_path, output_format, options):
                        success_count += 1
                        self.conversion_stats['downloads_completed'] += 1
                    if progress_callback:
                        progress = (i + 1) / total_files * 100
                        progress_callback(progress)
                except Exception as e:
                    self._update_stats(False)
                    continue
            return success_count > 0
        except Exception as e:
            if progress_callback:
                progress_callback(0)
            raise e

    def _update_stats(self, success):
        self.conversion_stats['total'] += 1
        if success:
            self.conversion_stats['successful'] += 1
        else:
            self.conversion_stats['failed'] += 1

    def get_conversion_stats_chart(self):
        fig, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(15, 4))

        # Pie chart for format distribution
        categories = list(self.supported_formats.keys())
        sizes = [len(self.supported_formats[cat]) for cat in categories]
        ax1.pie(sizes, labels=categories, autopct='%1.1f%%', startangle=90)
        ax1.set_title('Распределение поддерживаемых форматов')

        # Bar chart for conversion stats
        stats_labels = ['Успешно', 'Неудачно']
        stats_values = [self.conversion_stats['successful'], self.conversion_stats['failed']]
        colors = ['green', 'red']
        ax2.bar(stats_labels, stats_values, color=colors)
        ax2.set_title('Статистика конвертаций')
        ax2.set_ylabel('Количество')

        # Bar chart for download stats
        download_labels = ['Завершено', 'Неудачно']
        download_values = [self.conversion_stats['downloads_completed'], self.conversion_stats['failed']]
        ax3.bar(download_labels, download_values, color=['blue', 'red'])
        ax3.set_title('Статистика загрузок')
        ax3.set_ylabel('Количество')

        plt.tight_layout()
        buf = BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        img_data = base64.b64encode(buf.read()).decode('utf-8')
        plt.close(fig)
        return f"data:image/png;base64,{img_data}"

    def _convert_image(self, input_path, output_path, output_format, options):
        try:
            pil_format = self.pil_format_mapping.get(output_format, output_format.upper())
            with Image.open(input_path) as img:
                if options.get('resize_enabled', False):
                    width = options.get('width')
                    height = options.get('height')
                    if width and height:
                        if options.get('keep_aspect', True):
                            img.thumbnail((width, height), Image.Resampling.LANCZOS)
                        else:
                            img = img.resize((width, height), Image.Resampling.LANCZOS)
                if output_format in ['jpg', 'jpeg'] and img.mode in ['RGBA', 'LA', 'P']:
                    if img.mode == 'P' and 'transparency' in img.info:
                        img = img.convert('RGBA')
                    else:
                        img = img.convert('RGB')
                save_kwargs = {}
                if output_format in ['jpg', 'jpeg']:
                    save_kwargs['quality'] = options.get('quality', 95)
                elif output_format == 'png':
                    save_kwargs['optimize'] = True
                img.save(output_path, format=pil_format, **save_kwargs)
            return os.path.exists(output_path)
        except Exception as e:
            raise Exception(f"Image conversion error: {e}")

    def _convert_data(self, input_path, output_path, output_format):
        try:
            input_ext = os.path.splitext(input_path)[1].lower().lstrip('.')
            if input_ext in ['csv']:
                df = pd.read_csv(input_path)
            elif input_ext in ['xlsx', 'xls']:
                df = pd.read_excel(input_path)
            elif input_ext in ['json']:
                df = pd.read_json(input_path)
            elif input_ext in ['xml']:
                tree = ET.parse(input_path)
                data = [{child.tag: child.text for child in elem} for elem in tree.getroot()]
                df = pd.DataFrame(data)
            elif input_ext in ['yaml']:
                with open(input_path, 'r', encoding='utf-8') as f:
                    data = yaml.safe_load(f)
                df = pd.json_normalize(data)
            elif input_ext in ['h5']:
                with h5py.File(input_path, 'r') as f:
                    data = {key: f[key][()] for key in f.keys()}
                df = pd.DataFrame(data)
            elif input_ext in ['sql']:
                conn = sqlite3.connect(input_path)
                df = pd.read_sql_query("SELECT * FROM sqlite_master", conn)
                conn.close()
            elif input_ext in ['txt']:
                with open(input_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                if output_format == 'json':
                    data = {"content": content}
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)
                    return True
                elif output_format == 'csv':
                    lines = content.split('\n')
                    with open(output_path, 'w', encoding='utf-8', newline='') as f:
                        writer = csv.writer(f)
                        for line in lines:
                            if line.strip():
                                writer.writerow([line])
                    return True
                elif output_format in ['xlsx', 'xls']:
                    lines = content.split('\n')
                    df = pd.DataFrame(lines, columns=['Content'])
                    df.to_excel(output_path, index=False)
                    return True
                elif output_format == 'xml':
                    root = ET.Element("root")
                    item = ET.SubElement(root, "content")
                    item.text = content
                    tree = ET.ElementTree(root)
                    tree.write(output_path)
                    return True
                elif output_format == 'yaml':
                    data = {"content": content}
                    with open(output_path, 'w', encoding='utf-8') as f:
                        yaml.dump(data, f)
                    return True
                else:
                    return self._copy_file(input_path, output_path)
            else:
                raise ValueError(f"Unsupported input format: {input_ext}")
            if output_format == 'csv':
                df.to_csv(output_path, index=False, encoding='utf-8')
            elif output_format == 'json':
                df.to_json(output_path, orient='records', indent=2, force_ascii=False)
            elif output_format in ['xlsx', 'xls']:
                df.to_excel(output_path, index=False)
            elif output_format == 'txt':
                df.to_csv(output_path, index=False, sep='\t', encoding='utf-8')
            elif output_format == 'xml':
                data = df.to_dict('records')
                root = ET.Element("root")
                for record in data:
                    item = ET.SubElement(root, "item")
                    for key, value in record.items():
                        child = ET.SubElement(item, key)
                        child.text = str(value)
                ET.ElementTree(root).write(output_path)
            elif output_format == 'yaml':
                data = df.to_dict('records')
                with open(output_path, 'w', encoding='utf-8') as f:
                    yaml.dump(data, f)
            elif output_format == 'h5':
                with h5py.File(output_path, 'w') as f:
                    for column in df.columns:
                        f.create_dataset(column, data=df[column].values)
            elif output_format == 'sql':
                conn = sqlite3.connect(output_path)
                df.to_sql('data', conn, index=False, if_exists='replace')
                conn.close()
            return True
        except Exception as e:
            raise Exception(f"Data conversion error: {e}")

    def _convert_document(self, input_path, output_path, output_format):
        try:
            input_ext = os.path.splitext(input_path)[1].lower().lstrip('.')
            if input_ext == 'pdf' and output_format == 'txt':
                with fitz.open(input_path) as doc:
                    text = ""
                    for page in doc:
                        text += page.get_text()
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return True
            elif input_ext == 'pdf' and output_format in ['doc', 'docx']:
                text = ""
                with fitz.open(input_path) as doc:
                    for page in doc:
                        text += page.get_text()
                doc = Document()
                doc.add_paragraph(text)
                doc.save(output_path)
                return True
            elif input_ext in ['doc', 'docx'] and output_format == 'pdf':
                pythoncom.CoInitialize()
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(input_path)
                doc.SaveAs(output_path, FileFormat=17)
                doc.Close()
                word.Quit()
                return True
            elif input_ext in ['doc', 'docx'] and output_format == 'txt':
                doc = Document(input_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return True
            elif input_ext == 'md' and output_format == 'html':
                with open(input_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                html_content = markdown.markdown(md_content)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return True
            elif input_ext == 'html' and output_format == 'md':
                with open(input_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                md_content = html2text.html2text(html_content)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                return True
            elif input_ext == 'txt' and output_format == 'docx':
                with open(input_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                doc = Document()
                doc.add_paragraph(text)
                doc.save(output_path)
                return True
            else:
                return self._copy_file(input_path, output_path)
        except Exception as e:
            raise Exception(f"Document conversion error: {e}")

    def _convert_audio(self, input_path, output_path, output_format):
        try:
            audio = AudioSegment.from_file(input_path)
            audio.export(output_path, format=output_format)
            return True
        except Exception as e:
            raise Exception(f"Audio conversion error: {e}")

    def _convert_video(self, input_path, output_path, output_format):
        try:
            return self._copy_file(input_path, output_path)
        except Exception as e:
            raise Exception(f"Video conversion error: {e}")

    def _convert_archive(self, input_path, output_path, output_format):
        try:
            temp_dir = tempfile.mkdtemp()
            if input_path.endswith('.zip'):
                with zipfile.ZipFile(input_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
            elif input_path.endswith('.tar') or input_path.endswith('.tar.gz') or input_path.endswith('.tar.bz2') or input_path.endswith('.tar.xz'):
                with tarfile.open(input_path, 'r:*') as tar_ref:
                    tar_ref.extractall(temp_dir)
            elif input_path.endswith('.rar'):
                with rarfile.RarFile(input_path, 'r') as rar_ref:
                    rar_ref.extractall(temp_dir)
            elif input_path.endswith('.7z'):
                with py7zr.SevenZipFile(input_path, 'r') as seven_ref:
                    seven_ref.extractall(temp_dir)
            else:
                return self._copy_file(input_path, output_path)
            if output_format == 'zip':
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                    for root, _, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, temp_dir)
                            zip_ref.write(file_path, arcname)
            elif output_format in ['tar', 'gz', 'bz2', 'xz']:
                mode = 'w:gz' if output_format == 'gz' else 'w:bz2' if output_format == 'bz2' else 'w:xz' if output_format == 'xz' else 'w'
                with tarfile.open(output_path, mode) as tar_ref:
                    for root, _, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, temp_dir)
                            tar_ref.add(file_path, arcname)
            elif output_format == '7z':
                with py7zr.SevenZipFile(output_path, 'w') as seven_ref:
                    for root, _, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, temp_dir)
                            seven_ref.write(file_path, arcname)
            elif output_format == 'rar':
                return self._copy_file(input_path, output_path)
            shutil.rmtree(temp_dir)
            return True
        except Exception as e:
            raise Exception(f"Archive conversion error: {e}")

    def get_video_preview(self, input_path, frame_time=1.0, output_image_path=None):
        cap = cv2.VideoCapture(input_path)
        fps = cap.get(cv2.CAP_PROP_FPS)
        if fps <= 0:
            fps = 25
        frame_number = int(fps * frame_time)
        cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
        ret, frame = cap.read()
        cap.release()
        if not ret:
            raise Exception("Unable to extract frame from video")
        if output_image_path:
            cv2.imwrite(output_image_path, frame)
            return output_image_path
        return frame

    def _copy_file(self, input_path, output_path):
        try:
            shutil.copy2(input_path, output_path)
            return True
        except Exception as e:
            raise Exception(f"File copy error: {e}")

    def get_file_info(self, file_path):
        if not os.path.exists(file_path):
            return None
        try:
            info = {
                'name': os.path.basename(file_path),
                'size': os.path.getsize(file_path),
                'format': os.path.splitext(file_path)[1].upper().lstrip('.'),
                'created': os.path.getctime(file_path)
            }
            if info['format'].lower() in self.supported_formats['images']:
                try:
                    with Image.open(file_path) as img:
                        info['width'] = img.width
                        info['height'] = img.height
                        info['mode'] = img.mode
                except:
                    pass
            elif info['format'].lower() in self.supported_formats['audio']:
                try:
                    if info['format'].lower() == 'wav':
                        with contextlib.closing(wave.open(file_path, 'r')) as f:
                            info['channels'] = f.getnchannels()
                            info['sample_rate'] = f.getframerate()
                            info['duration'] = f.getnframes() / float(f.getframerate())
                    else:
                        audio = AudioSegment.from_file(file_path)
                        info['duration'] = len(audio) / 1000
                        info['channels'] = audio.channels
                        info['sample_rate'] = audio.frame_rate
                except:
                    pass
            elif info['format'].lower() in self.supported_formats['video']:
                info['type'] = 'video'
                info['duration'] = 'Unknown'
            elif info['format'].lower() in self.supported_formats['archives']:
                info['type'] = 'archive'
            elif info['format'].lower() in self.supported_formats['data']:
                info['type'] = 'data'
            elif info['format'].lower() in self.supported_formats['documents']:
                info['type'] = 'document'
            elif info['format'].lower() in self.supported_formats['code']:
                info['type'] = 'code'
            else:
                info['type'] = 'other'
            return info
        except Exception as e:
            raise Exception(f"Error retrieving file info: {e}")